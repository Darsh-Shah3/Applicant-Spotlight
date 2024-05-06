import streamlit as st 

import docx
from docx.shared import Pt
from docx import Document

from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING

from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from dotenv import load_dotenv
load_dotenv()
import os
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

import llama_index
from llama_index.core import Document, VectorStoreIndex, ServiceContext, SimpleDirectoryReader
from llama_index.core import Settings
from llama_index.llms.openai import OpenAI
import PyPDF2
import docx2txt

st.title('Resume Enhancer')

title_query = '''you have resume data from that extract the name of the person,experience,and job title.
In following format :

Name of the person - job title ([number] years of Experience).'''


summary_query = '''extract key details from provided text in order to write a summary which should formal and point wise, short sentances, and formal way.make the detail like directly writen in resume. 
below some predefined word format are there, so output needed in proper to make sure:

Title (like: Summary)
[points]

***output should be in proper format***                          

Summary
years of experience as a developer...
...needed a summary containing only the essential related points.
...

'''

skill_query = '''extract the only skills from the text.details should be in formal and simple.  
below some predefined word format are there, so output needed in below proper to make sure:

Title (like :Skills)
[skill name or category]

Examples : [Programming Languages]: [Python, Angular, ...]
 
***output should be in proper format***

Skills

Programming Languages: Python, Angular, JavaScript, SQL
Framework: Flask
Tools: Kubernetes, Jupiter Notebooks, Kafta, Docker
...for all skill

'''

project_query = '''you have details from that you have to extract all the projects related all detailed information.  
below some predefined format are there, so output needed in below proper format to make sure:

[only project title] 
Role : [role of the candidate in that project.if not mentioned then write " Devloper"] 
Description : [write the full description of the project]
Technology : [Extract the key technologies used in the provided project details.]
 
***output should be in proper format***
*** role,description,technology add for each project.if not then create detail based on job title from content

Example :
1. Mining Site Report generation
Role: Team Lead 
description : A definite report of ...
Technology : python,ruby,java

...
...all projects should be there

'''

role_query = '''

Roles and Responsibilities
List out all points starting key roles and responsibilities related to the [job title] position, based on the background details provided.

Example point:
 Managed software development of ...

If any details are missing from the provided content, create Roles and Responsibilities infomation based on that person background to complete the requested resume format. Focus on tailoring details to a specific job title.

***output should be in proper format***

Roles and Responsibilities
- point
- point
...

'''

def updated_skill(skill_text):
    from openai import OpenAI
    client = OpenAI()
    response = client.chat.completions.create(model="gpt-4",messages=[{"role": "system", "content": "You are a intelligent coder assistant."},
    {"role": "user", "content": f"understand this skill details : {skill_text}.understand skills details and based on that make two columns for main category and right side subcategory needed.output two column sperate wtih '|'."}])
    return response.choices[0].message.content

from llama_index.core import Settings
from llama_index.core import ServiceContext, PromptHelper

def first_query_engine():
    documents = SimpleDirectoryReader("C:\\Python_project\\Darsh_project_1\\HR_Gen_1\\txt_file").load_data() 
    from openai import OpenAI
    llm = OpenAI()
    service_context = ServiceContext.from_defaults() 
    index1 = VectorStoreIndex.from_documents(documents, service_context=service_context)
    query_engine1 = index1.as_query_engine()
    return query_engine1


def engine_query(query1,engine):
    response = engine.query(query1)
    return response.response



def create_summary_docx(title,summary,skill,role,resume_format):  
    filename = '/home/kp/Desktop/HR_Gen_1/logo/logo.png'
    width = Cm(5.44)
    height = Cm(0.69)
    color = RGBColor(204,64,37)
        
    doc = docx.Document() 
    section = doc.sections[0]
    header = section.header
    header.add_paragraph().add_run().add_picture(filename, width=width, height=height)
    header.top_margin = Pt(1 * 28.35)
    header.add_paragraph()


    # started title 
    doc.add_paragraph()
    title_paragraph = doc.add_paragraph(title) 
    title_paragraph.paragraph_format.line_spacing = 0.8
    title_run = title_paragraph.runs[0]
    title_run.bold = False
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(102,101,101) 
    title_run.font.name = 'Arial MT'
        
    #ended title

    # line
    def add_horizontal_line(paragraph, rgb_color, thickness_pt):
        pBdr = OxmlElement('w:pBdr')
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'single')
        bottom_border.set(qn('w:color'), rgb_color)
        pBdr.append(bottom_border)
        paragraph._element.get_or_add_pPr().append(pBdr)

    line_paragraph = doc.add_paragraph()
    add_horizontal_line(line_paragraph, '585353', 6)   
    line_paragraph.paragraph_format.line_spacing = 0.1
    doc.add_paragraph()

    #summray
    summary = summary.replace('-','')
    lines = summary.split('\n')
    title1 = doc.add_paragraph('Summary ')
    run = title1.runs[0]
    run.font.name = 'calibri'
    run.font.size = Pt(15)
    run.bold = True
    run.font.color.rgb = color   
    title1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    lines = lines[1:]  
    for point in lines:
        point = point.strip()
        if point.strip() != '':
            p = doc.add_paragraph(point, style='List Bullet')
            p.style.font.size = Pt(17)   
            run = p.runs[0]
            run.font.size = Pt(11)
            run.font.name = 'calibri'
            p.paragraph_format.line_spacing = 1.80

    #skills
    title1 = doc.add_paragraph("Skills")
    run = title1.runs[0]
    run.bold = True
    run.font.name = 'calibri'
    run.font.size = Pt(15)
    run.font.color.rgb = color 
    title1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    data = []
    lines = skill.strip().split('\n')[3:]  

    for line in lines:
        if line != '' :
            if '|' in line:
                main_category, sub_category_str = line.split('|')
                main_category = main_category.strip()
                sub_categories = [sub.strip() for sub in sub_category_str.split(',')]
                data.append((main_category, sub_categories))

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Main Category'
    header_cells[1].text = 'Sub Category'

    for main_category, sub_categories in data:
        row_cells = table.add_row().cells
        row_cells[0].text = main_category
        row_cells[1].text = sub_categories[0]

        for sub_category in sub_categories[1:]:
            row_cells = table.add_row().cells
            row_cells[1].text = sub_category

        if len(sub_categories) > 1:  
            a = table.cell(len(table.rows) - len(sub_categories), 0)
            b = table.cell(len(table.rows) - 1, 0)
            a.merge(b)

    for row in table.rows:
       for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'calibri'  # Set font to Arial
                    run.font.size = Pt(12)

    doc.add_paragraph() 
        
    #role
    role = role.replace('-','')
    lines = role.split('\n')
    title1 = doc.add_paragraph('Roles and Responsibilities')
    run = title1.runs[0]
    run.bold = True
    run.font.name = 'calibri'
    run.font.size = Pt(15)
    run.font.color.rgb = color   
    title1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    line = lines[1:]  
        
    for p1 in line:
        if p1.strip() != '':
            p1 = p1.strip()
            p = doc.add_paragraph(p1, style='List Bullet')
            p.style.font.size = Pt(18)   
            run = p.runs[0]
            run.font.size = Pt(11)
            run.font.name = 'calibri'
            p.paragraph_format.line_spacing = 1.5
    doc.add_paragraph()
    return doc


def add_project(doc,project_name, role, description,tech):
    
    # Add project name
    project_heading = doc.add_paragraph(f'{project_name}')
    project_heading.runs[0].bold = True
    project_heading.runs[0].font.size = Pt(12)
    project_heading.runs[0].font.name = 'calibri'
    project_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add role
    role = role.split(': ', 1)[1]
    role_heading = doc.add_paragraph('Role: ')
    run = role_heading.runs[0]
    run.font.name = 'calibri'
    run.font.size = Pt(12)
    run.bold = True 
    title_run = role_heading.add_run(role)
    title_run.font.name = 'calibri'
    title_run.font.size = Pt(11)

    # Add description
    role = description.split(': ', 1)[1]
    description_heading = doc.add_paragraph('Descriptions: ')
    run = description_heading.runs[0]
    run.font.name = 'calibri'
    run.font.size = Pt(12)
    run.bold = True 
    description_heading = description_heading.add_run(role)
    description_heading.font.name = 'calibri'
    description_heading.font.size = Pt(11)  
    
    # Add technology
    role = tech.split(': ', 1)[1]
    tech_heading = doc.add_paragraph('Technology: ')
    run = tech_heading.runs[0]
    run.font.name = 'calibri'
    run.font.size = Pt(12)
    run.bold = True 
    tech_heading = tech_heading.add_run(role)
    tech_heading.font.name = 'calibri'
    tech_heading.font.size = Pt(11) 
    doc.add_paragraph() 

def parse_projects(projects_list):
    for idx, project in enumerate(projects_list,start = 1):
        parts = project.split('\n')
        
        if len(parts) > 3:
            title, role, description, tech = parts  
            yield title, role, description, tech 
        elif len(parts) == 3:
            title, role, description = parts  
            tech = "None"
            yield title, role, description,tech 
        else:
            print(f"Skipping project {idx} due to unexpected format.")

def df():
    doc = Document()

    # Function to add a horizontal line
    def add_horizontal_line():
        line_paragraph = doc.add_paragraph()
        pBdr = OxmlElement('w:pBdr')
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'single')
        pBdr.append(bottom_border)
        line_paragraph._element.get_or_add_pPr().append(pBdr)

    info_query = """
    Based on the provided resume data, can you summarize the following personal information:

    * Name of the person
    * Phone no.
    * Email id
    * Location
    * LinkedIn id (if available)
    * Github id (if available)
    * Job title
    * Years of experience (if available)

    Give me answer in only dictionary format"""

    info_1 = engine_query(info_query,engine)
    info = eval(info_1)


    # Create the first heading
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = heading.add_run(info['Name'])
    run1.bold = True
    run1.font.size = Pt(35)
    run1.font.color.rgb = RGBColor(0, 0, 0)

    # Add a line break
    heading.add_run("\n")

    run2 = heading.add_run(info['Job title'])
    run2.bold = False
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(128, 128, 128)  # Gray color (optional)
    add_horizontal_line()


    # Contact Information with horizontal line
    doc.add_paragraph('\n')  # Add a space before the section
    doc.add_heading('CONTACT INFO', level=2)

    contact_info = doc.add_paragraph()

    # list = []
    for i,j in info.items():
        contact_info_run = contact_info.add_run(j+'\n')
        contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_horizontal_line()


    About_me_query = '''extract key details from provided text in order to write 'About me' section which should formal and point wise, short sentances, and formal way.make the detail like directly writen in resume. 
    .Need a formal, point-wise 'About Me' section, resembling a resume.Needed only four-five lines'''

    About_me = engine_query(About_me_query,engine)


    section_title = doc.add_heading(level=2)
    section_title_run = section_title.add_run('\nABOUT ME\n')
    section_title_run.bold = True
    section_content = doc.add_paragraph(About_me)
    section_content.paragraph_format.space_after = Pt(12)
    doc.add_paragraph('\n')
    add_horizontal_line()

    # Function to format job positions and details
    def format_job_position(doc, position_title, company, time_frame, details):
        job_title = doc.add_paragraph()
        job_title_run = job_title.add_run(position_title)
        job_title_run.bold = True
        job_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        company_info = doc.add_paragraph()
        company_info_run = company_info.add_run(company + ' – ' + time_frame + '\n')
        company_info_run.italic = True
        company_info_run.font.size = Pt(10)
        company_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        details_para = doc.add_paragraph()
        for detail in details:
            details_para.add_run('• ' + detail + '\n')
        details_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        add_horizontal_line()

    # Employment History
    doc.add_paragraph('\n')  # Add a space before the section
    doc.add_heading('EMPLOYMENT HISTORY', level=2)

    # Web Developer Job
    format_job_position(
        doc,
        'Web Developer - Shacara Technologies',
        'Wuse 2, Abuja',
        'January 2020 - December 2020',
        [
            'Managed and archived quality documentation and participated in internal and external quality audits',
            'Resolved conflicts and negotiated agreements between parties in order to reach win-win solutions to disagreements and clarify misunderstandings'
        ]
    )

    # Design Job
    format_job_position(
        doc,
        'Design - Cyber Technologies',
        'Wuse 2, Abuja',
        'April 2020 - July 2020',
        [
            'Resolved conflicts and negotiated agreements between parties in order to reach win-win solutions to disagreements and clarify misunderstandings',
            'Managed and archived quality documentation and participated in internal and external quality audits'
        ]
    )

    edu_query = '''Extract education details, such as colleges, schools, and related information, from the provided files.
    Format the output as follows for each specific education entry:
    'City Royal Sec - Senior Secondary School Certificate (SSCE)',
    'Wuse 2, Abuja',
    'January 2020 - November 2020'
    Ensure the extracted details include school/college names, locations, and duration of study.'''

    edu = engine_query(edu_query,engine)  
    doc.add_paragraph()

    # Add other sections
    sections = [
        ('Education', edu),
        ('CERTIFICATION', 'Introduction to Java by Learn Quest on Coursera.\nInheritance and Data Structure in Java by PENN UNIVERSITY on Coursera.'),
        ('INTERESTS', 'Car Enthusiast\nTravelling\nMotorcycling\nDrawing\nListening Music'),
    ]

    for title, content in sections:
        section_title = doc.add_heading(level=2)
        section_title_run = section_title.add_run(title + '\n')
        section_title_run.bold = True
        section_content = doc.add_paragraph(content)
        section_content.paragraph_format.space_after = Pt(12)


    doc.add_paragraph('\n')  # Add a space before the section
    doc.add_heading('PROJECTS\n', level=2)

    project_query = '''you have details from that you have to extract all the projects related all detailed information.  
    below some predefined format are there, so output needed in below proper format to make sure:

    [only project title] 
    Role : [role of the candidate in that project.if not mentioned then write " Devloper"] 
    Description : [write the full description of the project]
    Technology : [Extract the key technologies used in the provided project details.]
    
    ***output should be in proper format***
    *** role,description,technology add for each project.if not then create detail based on job title from content

    Example :
    1. Mining Site Report generation
    Role: Team Lead 
    description : A definite report of ...
    Technology : python,ruby,java

    ...
    ...all projects should be there

    '''

    projects_string = engine_query(project_query,engine)

    # save_file_path = "None"
    last_2(doc,projects_string,save_file_path)

    # Save the document
    # doc.save('misha.docx')
    return True

def last_1(doc,projects_string,save_file_path):    
    # doc = Document()
    title = doc.add_paragraph('Projects')
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(15)
    run.font.name = 'calibri'
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT 
    doc.add_paragraph()

    projects_list = projects_string.split('\n\n')
    for title, role, description, tech in parse_projects(projects_list): 
        add_project(doc,title, role, description,tech)

    doc.save('output/final_resume.docx')


def last_2(doc,projects_string,save_file_path):    
    # doc = Document()
    title = doc.add_paragraph('Projects')
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(15)
    run.font.name = 'calibri'
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT 
    doc.add_paragraph()

    projects_list = projects_string.split('\n\n')
    for title, role, description, tech in parse_projects(projects_list): 
        add_project(doc,title, role, description,tech)

    doc.save('output/final_resume.docx')

import subprocess
def convert_to_pdf(save_file_path, output_dir):
    subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', save_file_path, '--outdir', output_dir])
    return True
 
save_file_path = os.path.join('output','final_resume.docx')
output_dir = os.path.join('output')

if 'allow1' not in st.session_state:
    st.session_state.allow1 = False

if 'allow2' not in st.session_state:
    st.session_state.allow2 = False

if 'directory_path' not in st.session_state:
    st.session_state.directory_path = False


with st.sidebar.form(key='pdf_upload_form', clear_on_submit=True):
    resume1 = st.file_uploader("Choose Vendor's Resume", type=['docx', 'pdf'], key='p1')
    radio = st.radio(label="Select Resume Format", options=['1st- Format', '2nd - Format'], horizontal=True)
    submit_button = st.form_submit_button(label='Generate Resume', on_click=None)
    
import tempfile 
import os
import aspose.words as aw
random_count='1' 
    
if submit_button:
    if resume1 is not None:
        with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
            tmp_file.write(resume1.getvalue())
            tmp_file_path = tmp_file.name 
            if resume1.type =="application/pdf":
                page = PyPDF2.PdfReader(tmp_file_path)
                extracted_text = '' 
                for i in range(len(page.pages)):
                    extracted_text += page.pages[i].extract_text()
            elif resume1.type =="application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                file_path = open(tmp_file_path, 'rb')
                extracted_text = docx2txt.process(file_path)
            else:
                st.error("Unsupported file format. Please upload a PDF or DOCX file.")

        #extra_code
        with open("C:\\Python_project\\Darsh_project_1\\HR_Gen_1\\txt_file\\text.txt",'w',encoding="utf-8") as f:
            f.write(extracted_text)

        engine = first_query_engine()
        if radio == '1st- Format':
            list1 = [title_query, summary_query,skill_query, role_query,project_query]
            list1_names = ['title', 'summary', 'skill', 'role','project']


            for i, name in enumerate(list1_names):
                if i != 2:
                    list1_names[i] = engine_query(list1[i], engine)
                    st.write(list1_names[i])
                else:
                    skill_text = engine_query(list1[i], engine)
                    list1_names[i] = updated_skill(skill_text)
                    st.write(list1_names[i])


            doc = create_summary_docx(list1_names[0], list1_names[1],list1_names[2],list1_names[3] ,resume_format = radio)

            # title = doc.add_paragraph('Projects')
            last_1(doc,list1_names[4], save_file_path)
         
        else:
            
            doc = df()
            # title = doc.add_paragraph('Projects')
            # last_2(doc,projects_string, save_file_path)
        
        # convert_to_pdf(save_file_path, output_dir)

        doc = aw.Document("C:\\Python_project\\Darsh_project_1\\HR_Gen_1\\output\\final_resume.docx")
        doc.save('C:\\Python_project\\Darsh_project_1\\HR_Gen_1\\output\\final_resume'+random_count+'.pdf')

        st.session_state.allow1 = True
        st.session_state.allow2 = True
        st.session_state.directory_path = True
    

file_path1 = 'C:\\Python_project\\Darsh_project_1\\HR_Gen_1\\output\\final_resume'+random_count+'.pdf'
file_path2 = 'C:\\Python_project\\Darsh_project_1\\HR_Gen_1\\output\\final_resume.docx'
if st.session_state.directory_path == True:
    if st.session_state.allow1 == True:
        with open(file_path1, "rb") as file:
            if st.download_button(label="Download Document PDF", data=file, file_name="document.pdf", mime="application/pdf"):
                st.session_state.allow1 = False

if st.session_state.directory_path == True:
    if st.session_state.allow2 == True:
        with open(file_path2, "rb") as file:
            if st.download_button(label="Download Document DOCX", data=file, file_name="document.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
                st.session_state.allow2 = False
 
   
